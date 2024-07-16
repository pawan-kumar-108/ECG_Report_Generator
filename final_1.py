import streamlit as st
from pathlib import Path
import google.generativeai as genai
from api_key import API_KEY
from docx import Document
from io import BytesIO

genai.configure(api_key = API_KEY)

generation_config = {
  "temperature": 0.5,
  "top_p": 0.95,
  "top_k": 64,
  "max_output_tokens": 8192,
  "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
  model_name="gemini-1.5-flash",
  generation_config=generation_config
)

system_prompt="""
You are the world's top cardiologist and performed 1000 ECGs accoss the world. You are known for giving the best report after analyzing the ECG. 
Make sure that you removing the points which is not available in the ecg paper(for example, {Gender}, {ID Number}, {Signature}, or any other information that you are not able to carve out.) You have to analyze the uploaded {ecg_data}, the waves and give the below details in a pointwise manner (skip those which are not available). Don't describe the disease, just focus on describing the following data in a clear and concise manner.
(and print them in same format, delete those which can't be deduced)


Patient Information-

    Name: {{name}}
    Age: {{age}}
    Gender: {{gender}}
    ID Number: {{id_number}}
    Date of ECG: {{date_of_ecg}}

Clinical Information:

    Reason for ECG: {{reason_for_ecg}}
    Relevant Medical History: {{medical_history}}
    Medications: {{medications}}

ECG Technical Details:

    ECG Machine Used: {{ecg_machine}}
    Lead Configuration: {{lead_configuration}}
    Calibration: {{calibration}}
    Recording Quality: {{recording_quality}}

ECG Findings:

    Rhythm and Rate :-

        Heart Rate: {{heart_rate}} bpm
        Rhythm: {{rhythm}}
        P Waves: {{p_waves}}
        PR Interval: {{pr_interval}} ms
        QRS Complex: {{qrs_complex}} ms
        QT/QTc Interval: {{qt_qtc_interval}} ms
        ST Segment: {{st_segment}}
        T Waves: {{t_waves}}

    Axis:-

        P Wave Axis: {{p_wave_axis}} degrees
        QRS Axis: {{qrs_axis}} degrees
        T Wave Axis: {{t_wave_axis}} degrees

    Conduction and Morphology:-

        Atrial Conduction: {{atrial_conduction}}
        Ventricular Conduction: {{ventricular_conduction}}
        QRS Morphology: {{qrs_morphology}}
        ST-T Changes: {{st_t_changes}}

Interpretation-

    Normal or Abnormal: {{interpretation}}
    Diagnosis/Findings: {{diagnosis}}

Conclusion and Recommendations -

    Summary: {{summary}}
    Recommendations: {{recommendations}}

Reporting Cardiologist - 

    Name: {{cardiologist_name}}
    Signature: {{signature}}
    Date of Report: {{date_of_report}}

Attachments -

    ECG Tracing: {{ecg_tracing}}

You don't have to ask anything else. just generate whatever you analyzed from the graph.
Make sure that the identation and alignment pattern remain the same. (no issue for those which are not available, just delete that whole line)

Here is in what way you could delete:
for example, Gender is not available, then remove 'Gender: {{gender}}'
and let's suppose, neither Name: {{cardiologist_name}}
    nor Signature: {{signature}}
    nor Date of Report: {{date_of_report}} info is available, then delete complete section i.e. 'Reporting Cardiologist - ', same goes for all the sections like, 'Clinical Information', and any other section or sub-section like 'Conduction and Morphology:-', remove it if there is nothing that could be shown
But make sure that the 'Conclusion and Recommendations -' and in 'Interpretation -', there must be the required analysis.
Never keep {{summary}} or  {{recommendations}} empty! It will include the analysis of the report.
Maybe at some point, it could not be possible to deduce the {P Wave Axis} or {ST Segment} or any other segment, in that segment delete that relevant title like 'ST Segment'
Recommendation involves the step the patient to take next, based on the summary which is given by you, i.e the world class doctor!

CAUTION: NO SEGMNET SHOULD BE LEFT EMPTY. IF EMPTY REMOVE IT FROM YOUR RESPONSE.

"""
st.set_page_config(page_title = "ECG Paper Reader")

st.title("Your Digital ECG Reader: ")
st.subheader("It can analyze the ECG paper to know all the details about your heart. ❤️")

ecg_file = st.file_uploader("Upload your ECG image", type=["jpg", "jpeg", "png"])

generate_button = st.button("Get ECG report")
if generate_button:
    ecg_data = ecg_file.getvalue()
    mime_type = ecg_file.type
    image_parts = {
        "mime_type": mime_type,
        "data": ecg_data
    }

    prompt_parts = {
        "parts": [
            image_parts,
            {"text": system_prompt}
        ]
    }

    response = model.generate_content(prompt_parts)
    #st.write(response.text)
    
    # Create a Word document
    doc = Document()
    doc.add_heading('ECG Report', 0)
    doc.add_paragraph(response.text)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Provide download button for the Word document
    st.download_button(
        label="Download ECG Report",
        data=buffer,
        file_name="ecg_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )